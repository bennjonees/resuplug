from flask import Flask, request, send_file, jsonify
from docx import Document
from io import BytesIO
import openai
import os
import logging

app = Flask(__name__)

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Set your OpenAI API key
openai.api_key = os.getenv("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")  # Use environment variable on Render
# Uncomment and paste your key here for local testing:
# openai.api_key = "sk-..."

def enhance_text(prompt, style="You're a resume expert. Rewrite this to sound professional and concise."):
    """Enhance text using OpenAI ChatGPT with error handling."""
    if not prompt:
        return ""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": style},
                {"role": "user", "content": prompt}
            ]
        )
        return response['choices'][0]['message']['content'].strip()
    except Exception as e:
        logger.error(f"OpenAI error: {e}")
        return prompt  # Fallback to original text

def replace_text_in_doc(doc, replacements):
    """Replace placeholders in paragraphs and tables of a Word document."""
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, val in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, val)
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)

@app.route('/generate-resume', methods=['POST'])
def generate_resume():
    """Generate a resume from a template using provided data."""
    # Validate input
    data = request.json
    if not data:
        return jsonify({"error": "No JSON data provided"}), 400

    # Load template
    template_name = data.get("template", "adelinepalmerston")
    template_path = f"templates/{template_name}.docx"
    if not os.path.exists(template_path):
        return jsonify({"error": f"Template not found at {template_path}"}), 400

    try:
        doc = Document(template_path)
    except Exception as e:
        logger.error(f"Error loading template: {e}")
        return jsonify({"error": "Failed to load template"}), 500

    # Enhance text fields
    summary = enhance_text(data.get("summary", ""))
    education = enhance_text(data.get("education", ""))
    interests = enhance_text(data.get("interests", ""))
    references = enhance_text(data.get("references", ""))

    # Handle skills as a list
    skills = data.get("skills", [])  # Expecting a list like ["Python", "Flask", ...]
    skill_placeholders = [
        "[Skill 1]", "[Skill 2]", "[Skill 3]", "[Skill 4]",
        "[Skill 5]", "[Skill 6]", "[Skill 7]", "[Skill 8]"
    ]
    skill_replacements = {placeholder: skill for placeholder, skill in zip(skill_placeholders, skills)}

    # Define replacements
    replacements = {
        "[Your Full Name]": data.get("full_name", ""),
        "[Your Email Address]": data.get("email", ""),
        "[Your Phone Number]": data.get("phone", ""),
        "[Your Address]": data.get("address", ""),
        "[Write a short summary about your experience and goals]": summary,
        "[Your Degree], [Field of Study] [Years Attended]": education,
        "[Interests]": interests,
        "[References]": references,
        **skill_replacements  # Merge skill replacements
    }

    # Replace placeholders in the document
    replace_text_in_doc(doc, replacements)

    # Save to BytesIO and return
    result = BytesIO()
    doc.save(result)
    result.seek(0)
    return send_file(result, as_attachment=True, download_name="Resume.docx")

if __name__ == '__main__':
    app.run(debug=True)
